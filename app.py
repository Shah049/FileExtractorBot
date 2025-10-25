import streamlit as st
import os
import requests
from PIL import Image
import pytesseract
from docx import Document
import pandas as pd
from bs4 import BeautifulSoup
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain_community.chains import ConversationalRetrievalChain
from langchain.schema import Document as LCDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
import tempfile
import time

# --- Page Configuration ---
st.set_page_config(
    page_title="DocuChat 🤖",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS for Styling ---
st.markdown("""
<style>
    /* Main app styling */
    .stApp {
        background-color: #f0f2f6;
    }
    /* Title styling */
    h1 {
        color: #1E3A8A; /* Dark Blue */
        text-align: center;
    }
    /* Subheader styling */
    .stMarkdown {
        text-align: center;
    }
    /* Sidebar styling */
    .css-1d391kg {
        background-color: #FFFFFF;
    }
    /* Button styling */
    .stButton>button {
        background-color: #2563EB; /* Bright Blue */
        color: white;
        border-radius: 8px;
        border: none;
    }
    .stButton>button:hover {
        background-color: #1D4ED8; /* Darker Blue on hover */
    }
    /* Chat message styling */
    .st-chat-message-user {
        background-color: #E0E7FF; /* Light Blue for user */
    }
    .st-chat-message-assistant {
        background-color: #F3F4F6; /* Light Gray for assistant */
    }
</style>
""", unsafe_allow_html=True)


# --- Caching for expensive models ---
@st.cache_resource
def load_embedding_model():
    """Loads the HuggingFace embedding model."""
    return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")


@st.cache_resource
@st.cache_resource
def load_llm(api_key):
    """Loads the Groq LLM model."""
    return ChatGroq(api_key=api_key, model_name="llama-3.3-70b-versatile", temperature=0.3)



# --- Helper Functions for Text Extraction ---
def extract_text_from_file(uploaded_file):
    """Extracts text from various file formats."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        file_path = tmp_file.name
    try:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            documents = loader.load()
        elif ext == ".txt":
            from langchain_community.document_loaders import TextLoader
            loader = TextLoader(file_path)
            documents = loader.load()
        elif ext == ".docx":
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text])
            documents = [LCDocument(page_content=text)]
        elif ext in ['.xlsx', '.xls']:
            df = pd.read_excel(file_path)
            text = df.to_string()
            documents = [LCDocument(page_content=text)]
        elif ext in ['.png', '.jpg', '.jpeg']:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            documents = [LCDocument(page_content=text)]
        else:
            st.error(f"Unsupported file type: {ext}")
            return None
    except Exception as e:
        st.error(f"Error processing file: {e}")
        return None
    finally:
        try:
            os.remove(file_path)
        except Exception as e:
            st.warning(f"Could not delete temporary file: {e}")
    return documents


def extract_text_from_link(url):
    """Extracts text content from a web page."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")
        for script_or_style in soup(['script', 'style', 'header', 'footer', 'nav', 'aside']):
            script_or_style.decompose()
        text = soup.get_text(separator='\n', strip=True)
        return [LCDocument(page_content=text, metadata={"source": url})]
    except Exception as e:
        st.error(f"Failed to extract content from link: {e}")
        return None


# --- Dashboard Functions ---
def get_document_analytics(documents, llm):
    """Generates a summary and analytics for the document."""
    full_text = " ".join([doc.page_content for doc in documents])
    word_count = len(full_text.split())
    char_count = len(full_text)
    reading_time = round(word_count / 200)  # Average reading speed

    try:
        summary_prompt = f"Provide a concise, 3-sentence summary of the following text:\n\n{full_text[:4000]}"
        summary = llm.invoke(summary_prompt).content

        topics_prompt = f"Extract the 5 most important keywords or topics from the following text:\n\n{full_text[:4000]}"
        topics_response = llm.invoke(topics_prompt).content
        topics = [topic.strip() for topic in topics_response.replace("-", "").split("\n") if topic.strip()]

    except Exception as e:
        st.error(f"Could not generate analytics: {e}")
        summary = "Could not generate summary."
        topics = []

    return {
        "word_count": word_count,
        "char_count": char_count,
        "reading_time": reading_time,
        "summary": summary,
        "topics": topics
    }


# --- Main App UI ---
st.title("📄 DocuChat: Chat with Your Documents & Links 🤖")
st.markdown(
    "Upload a document or paste a web link, and I'll answer your questions about it using the power of LLaMA 3.")
st.markdown("---")

# --- Sidebar for Configuration ---
with st.sidebar:
    st.header("⚙️ Configuration")
    try:
        groq_api_key = st.secrets["GROQ_API_KEY"]
        st.success("GROQ API key loaded successfully!")
    except (KeyError, FileNotFoundError):
        st.error("GROQ_API_KEY not found in st.secrets!")
        st.info("Please add your GROQ API key to a .streamlit/secrets.toml file.")
        st.stop()

    age_group = st.radio("Select an explanation style:", ("Child", "Young Person", "Adult"), index=2, horizontal=True)
    st.markdown("---")
    st.markdown(
        """
        **How it works:**
        1.  Upload a document or paste a URL.
        2.  The content is extracted and vectorized.
        3.  Ask questions and get answers from the document!
        """
    )
    st.markdown("---")
    st.markdown("Built with ❤️ by combining LangChain, Groq, and Streamlit.")

# --- Session State Initialization ---
if "qa_chain" not in st.session_state:
    st.session_state.qa_chain = None
if "messages" not in st.session_state:
    st.session_state.messages = []
if "processing_done" not in st.session_state:
    st.session_state.processing_done = False
if "analytics" not in st.session_state:
    st.session_state.analytics = None


# --- Input Handling (File and Link) ---
def process_documents(documents, llm):
    if not documents:
        st.warning("Could not extract any content. Please try a different source.")
        return

    with st.spinner("Embedding text and building the chatbot... This is the magic part!"):
        try:
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            texts = text_splitter.split_documents(documents)
            if not texts:
                st.error("No valid text chunks extracted from the document. Try a different file or URL.")
                return

            embeddings = load_embedding_model()
            db = FAISS.from_documents(texts, embeddings)
            retriever = db.as_retriever()

            st.session_state.qa_chain = ConversationalRetrievalChain.from_llm(llm, retriever)
            st.session_state.processing_done = True
            st.session_state.messages = []  # Reset chat history
            st.success("✅ Chatbot is ready! Ask your questions in the 'Chat' tab.")

            with st.spinner("Generating document analytics..."):
                st.session_state.analytics = get_document_analytics(documents, llm)
            st.rerun()  # Rerun to switch to the chat/dashboard view cleanly

        except Exception as e:
            st.error(f"Error building chatbot: {e}")


# Main content area
if not st.session_state.processing_done:
    input_tabs = st.tabs(["📁 Upload a File", "🔗 Paste a Link"])
    with input_tabs[0]:
        uploaded_file = st.file_uploader("Choose a file",
                                         type=['pdf', 'txt', 'docx', 'xlsx', 'xls', 'png', 'jpg', 'jpeg'])
        if st.button("Process File", key="process_file"):
            if uploaded_file:
                with st.spinner("Reading and processing file..."):
                    llm = load_llm(groq_api_key)
                    docs = extract_text_from_file(uploaded_file)
                    process_documents(docs, llm)
            else:
                st.warning("Please upload a file first.")

    with input_tabs[1]:
        url_input = st.text_input("Enter the URL of a webpage:")
        if st.button("Process Link", key="process_link"):
            if url_input:
                with st.spinner("Fetching and processing webpage content..."):
                    llm = load_llm(groq_api_key)
                    docs = extract_text_from_link(url_input)
                    process_documents(docs, llm)
            else:
                st.warning("Please enter a URL first.")
else:
    # --- Chat and Dashboard Interface ---
    chat_tabs = st.tabs(["💬 Chat", "📊 Dashboard"])

    with chat_tabs[0]:
        # Display chat messages from history
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        # Accept user input
        if prompt := st.chat_input("Ask a question about your content..."):
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": prompt})

            # Generate and add assistant response
            with st.spinner("Thinking..."):
                chat_history_tuples = []
                for i in range(0, len(st.session_state.messages) - 2, 2):  # Go up to the second to last message
                    if (st.session_state.messages[i]["role"] == "user" and
                            i + 1 < len(st.session_state.messages) and
                            st.session_state.messages[i + 1]["role"] == "assistant"):
                        chat_history_tuples.append(
                            (st.session_state.messages[i]["content"], st.session_state.messages[i + 1]["content"])
                        )

                formatted_query = f"Based on the provided document, explain this to a {age_group.lower()}: {prompt}"
                try:
                    result = st.session_state.qa_chain.invoke(
                        {"question": formatted_query, "chat_history": chat_history_tuples}
                    )
                    answer = result["answer"]
                    st.session_state.messages.append({"role": "assistant", "content": answer})
                except Exception as e:
                    st.session_state.messages.append(
                        {"role": "assistant", "content": f"Error generating response: {e}"})

            # Rerun to display the new messages in the correct order
            st.rerun()

    with chat_tabs[1]:
        if st.session_state.analytics:
            st.header("Document Analytics")
            analytics = st.session_state.analytics

            col1, col2, col3 = st.columns(3)
            col1.metric("Word Count", f"{analytics['word_count']:,}")
            col2.metric("Character Count", f"{analytics['char_count']:,}")
            col3.metric("Est. Reading Time", f"{analytics['reading_time']} min")

            st.subheader("📄 Content Summary")
            st.info(analytics['summary'])

            st.subheader("🔑 Key Topics")
            for topic in analytics['topics']:
                st.markdown(f"- **{topic}**")
        else:
            st.info("Analytics are being generated...")
